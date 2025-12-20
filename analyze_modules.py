#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Module Performance Analyzer - First Principles Optimization
分析各模組表現，識別誤報源，提出優化建議
"""
import json
import pandas as pd
import numpy as np
from pathlib import Path
from collections import defaultdict
from datetime import datetime
import sys
import io

BASE_DIR = Path(r"C:\Users\s_robby518\Documents\trae_projects\ai testing\tiktok_labeler\tiktok videos download")
DATA_DIR = BASE_DIR / "data"
OUTPUT_DIR = Path("output")
REPORT_DIR = BASE_DIR / "data" / "report"

# Ensure report directory exists
REPORT_DIR.mkdir(parents=True, exist_ok=True)

def load_training_data():
    """Load merged training dataset"""
    training_path = DATA_DIR / "training_dataset_full.xlsx"
    if not training_path.exists():
        print(f"[ERROR] Training dataset not found: {training_path}")
        return None

    df = pd.read_excel(training_path)
    print(f"[OK] Loaded {len(df)} training samples")
    return df

def analyze_false_positives(df):
    """Analyze false positives - AI says AI but human says REAL"""
    fp_data = df[df['False_Positive'] == True].copy()

    if len(fp_data) == 0:
        print("\n[OK] No false positives found!")
        return None

    print(f"\n{'='*80}")
    print(f"FALSE POSITIVE ANALYSIS (AI說AI但人說REAL)")
    print(f"{'='*80}")
    print(f"Total false positives: {len(fp_data)}/{len(df)} ({len(fp_data)/len(df)*100:.1f}%)")

    # Module score analysis for false positives
    module_cols = [
        'metadata_extractor', 'frequency_analyzer', 'texture_noise_detector',
        'model_fingerprint_detector', 'lighting_geometry_checker',
        'heartbeat_detector', 'blink_dynamics_analyzer', 'av_sync_verifier',
        'text_fingerprinting', 'semantic_stylometry',
        'sensor_noise_authenticator', 'physics_violation_detector'
    ]

    # Calculate average scores for false positives vs correct predictions
    fp_means = fp_data[module_cols].mean()
    correct_real = df[(df['Human_Label'] == 'REAL') & (df['AI_Says_REAL'] == True)]
    correct_means = correct_real[module_cols].mean() if len(correct_real) > 0 else pd.Series()

    print(f"\nModule Score Comparison (False Positives vs Correct REAL):")
    print(f"{'Module':<30} {'FP Avg':<12} {'Correct Avg':<12} {'Diff':<12} {'Issue?'}")
    print(f"{'-'*80}")

    problem_modules = []
    for module in module_cols:
        fp_avg = fp_means[module]
        correct_avg = correct_means[module] if len(correct_means) > 0 else 0
        diff = fp_avg - correct_avg

        # Flag modules with high scores in false positives
        is_problem = fp_avg > 50 or diff > 30
        flag = "*** PROBLEM" if is_problem else ""

        if is_problem:
            problem_modules.append({
                'module': module,
                'fp_avg': fp_avg,
                'correct_avg': correct_avg,
                'diff': diff
            })

        print(f"{module:<30} {fp_avg:>10.1f}  {correct_avg:>10.1f}  {diff:>+10.1f}  {flag}")

    return problem_modules, fp_data

def analyze_false_negatives(df):
    """Analyze false negatives - AI says REAL but human says AI"""
    fn_data = df[df['False_Negative'] == True].copy()

    if len(fn_data) == 0:
        print("\n[OK] No false negatives found!")
        return None

    print(f"\n{'='*80}")
    print(f"FALSE NEGATIVE ANALYSIS (AI說REAL但人說AI)")
    print(f"{'='*80}")
    print(f"Total false negatives: {len(fn_data)}/{len(df)} ({len(fn_data)/len(df)*100:.1f}%)")

    # Module score analysis
    module_cols = [
        'metadata_extractor', 'frequency_analyzer', 'texture_noise_detector',
        'model_fingerprint_detector', 'lighting_geometry_checker',
        'heartbeat_detector', 'blink_dynamics_analyzer', 'av_sync_verifier',
        'text_fingerprinting', 'semantic_stylometry',
        'sensor_noise_authenticator', 'physics_violation_detector'
    ]

    fn_means = fn_data[module_cols].mean()
    correct_ai = df[(df['Human_Label'] == 'AI') & (df['AI_Says_AI'] == True)]
    correct_means = correct_ai[module_cols].mean() if len(correct_ai) > 0 else pd.Series()

    print(f"\nModule Score Comparison (False Negatives vs Correct AI):")
    print(f"{'Module':<30} {'FN Avg':<12} {'Correct Avg':<12} {'Diff':<12} {'Issue?'}")
    print(f"{'-'*80}")

    weak_modules = []
    for module in module_cols:
        fn_avg = fn_means[module]
        correct_avg = correct_means[module] if len(correct_means) > 0 else 0
        diff = fn_avg - correct_avg

        # Flag modules with low scores in false negatives (should be high for AI)
        is_weak = fn_avg < 40 or diff < -20
        flag = "*** WEAK" if is_weak else ""

        if is_weak:
            weak_modules.append({
                'module': module,
                'fn_avg': fn_avg,
                'correct_avg': correct_avg,
                'diff': diff
            })

        print(f"{module:<30} {fn_avg:>10.1f}  {correct_avg:>10.1f}  {diff:>+10.1f}  {flag}")

    return weak_modules, fn_data

def generate_optimization_recommendations(problem_modules, weak_modules, df):
    """Generate first-principles optimization recommendations"""
    print(f"\n{'='*80}")
    print("OPTIMIZATION RECOMMENDATIONS (第一性原理)")
    print(f"{'='*80}")

    recommendations = []

    # Analyze video characteristics of false positives
    fp_data = df[df['False_Positive'] == True]
    if len(fp_data) > 0:
        print(f"\n1. FALSE POSITIVE CHARACTERISTICS:")
        print(f"   Average AI_Probability: {fp_data['AI_Probability'].mean():.1f}%")
        print(f"   Average Bitrate: {fp_data['Bitrate'].mean()/1000000:.2f} Mbps")
        print(f"   Average Face Presence: {fp_data['Face_Presence'].mean():.1f}%")
        print(f"   Phone videos: {fp_data['Is_Phone'].sum()}/{len(fp_data)}")

        # Check if low bitrate causing issues
        if fp_data['Bitrate'].mean() < 2000000:
            recommendations.append({
                'type': 'THRESHOLD',
                'target': 'Low bitrate protection',
                'action': 'Reduce frequency_analyzer, texture_noise_detector weights for bitrate < 2 Mbps',
                'reason': 'Low bitrate videos naturally have compression artifacts'
            })

    # Problem modules (over-flagging)
    if problem_modules:
        print(f"\n2. OVER-FLAGGING MODULES (需要降低敏感度):")
        for pm in problem_modules:
            print(f"   - {pm['module']}: {pm['fp_avg']:.1f} in FP vs {pm['correct_avg']:.1f} in correct")

            recommendations.append({
                'type': 'MODULE_PARAM',
                'target': pm['module'],
                'action': f"Reduce sensitivity or weight by {min(50, pm['diff']/pm['fp_avg']*100):.0f}%",
                'reason': f"Score {pm['diff']:.1f} points higher in false positives"
            })

    # Weak modules (under-detecting)
    if weak_modules:
        print(f"\n3. UNDER-DETECTING MODULES (需要提高敏感度):")
        for wm in weak_modules:
            print(f"   - {wm['module']}: {wm['fn_avg']:.1f} in FN vs {wm['correct_avg']:.1f} in correct")

            recommendations.append({
                'type': 'MODULE_PARAM',
                'target': wm['module'],
                'action': f"Increase sensitivity or weight by {min(50, abs(wm['diff'])/wm['correct_avg']*100):.0f}%",
                'reason': f"Score {abs(wm['diff']):.1f} points lower in false negatives"
            })

    # Overall accuracy
    accuracy = (df['Correct'].sum() / len(df)) * 100
    print(f"\n4. CURRENT ACCURACY: {accuracy:.1f}%")

    if accuracy < 90:
        recommendations.append({
            'type': 'THRESHOLD',
            'target': 'Global thresholds',
            'action': 'Consider adjusting SAFE/GRAY/KILL thresholds based on distribution',
            'reason': f'Current accuracy {accuracy:.1f}% below target 90%'
        })

    return recommendations

def export_recommendations(recommendations, output_path):
    """Export recommendations to file"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# Module Optimization Recommendations\n")
        f.write("# Generated by First Principles Analysis\n\n")

        for i, rec in enumerate(recommendations, 1):
            f.write(f"{i}. Type: {rec['type']}\n")
            f.write(f"   Target: {rec['target']}\n")
            f.write(f"   Action: {rec['action']}\n")
            f.write(f"   Reason: {rec['reason']}\n\n")

    print(f"\n[OK] Recommendations exported to: {output_path}")

def generate_word_report(report_content, timestamp_str):
    """Generate Word document report"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading('Module Performance Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Timestamp
        timestamp_para = doc.add_paragraph(f'Generated: {timestamp_str}')
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Empty line

        # Add report content
        for line in report_content.split('\n'):
            if line.strip().startswith('==='):
                # Section header
                doc.add_heading(line.strip('= '), level=2)
            elif line.strip().startswith('---'):
                # Subsection separator
                doc.add_paragraph('_' * 80)
            else:
                # Regular content
                doc.add_paragraph(line)

        # Save
        filename = f"module_analysis_report_{timestamp_str}.docx"
        filepath = REPORT_DIR / filename
        doc.save(str(filepath))
        print(f"\n[OK] Word report saved: {filepath}")
        return filepath

    except ImportError:
        print("\n[WARN] python-docx not installed. Skipping Word report.")
        print("       Install with: pip install python-docx")
        return None
    except Exception as e:
        print(f"\n[ERROR] Failed to generate Word report: {e}")
        return None

def generate_pdf_report(report_content, timestamp_str):
    """Generate PDF report"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        filename = f"module_analysis_report_{timestamp_str}.pdf"
        filepath = REPORT_DIR / filename

        # Create PDF
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # Styles
        styles = getSampleStyleSheet()

        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#2C3E50',
            spaceAfter=30,
            alignment=TA_CENTER
        )

        # Heading style
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor='#34495E',
            spaceAfter=12,
            spaceBefore=12
        )

        # Normal style
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            leading=14
        )

        # Build content
        story = []

        # Title
        story.append(Paragraph("Module Performance Analysis Report", title_style))
        story.append(Paragraph(f"Generated: {timestamp_str}", normal_style))
        story.append(Spacer(1, 20))

        # Process content
        lines = report_content.split('\n')
        for line in lines:
            if line.strip().startswith('==='):
                # Section header
                story.append(Spacer(1, 12))
                story.append(Paragraph(line.strip('= '), heading_style))
            elif line.strip().startswith('---'):
                # Separator
                story.append(Spacer(1, 6))
            elif line.strip():
                # Regular content - escape HTML characters
                escaped_line = line.replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(escaped_line, normal_style))
            else:
                story.append(Spacer(1, 6))

        # Build PDF
        doc.build(story)
        print(f"[OK] PDF report saved: {filepath}")
        return filepath

    except ImportError:
        print("\n[WARN] reportlab not installed. Skipping PDF report.")
        print("       Install with: pip install reportlab")
        return None
    except Exception as e:
        print(f"\n[ERROR] Failed to generate PDF report: {e}")
        return None

def capture_output(func):
    """Decorator to capture function output"""
    def wrapper(*args, **kwargs):
        # Capture stdout
        old_stdout = sys.stdout
        sys.stdout = captured_output = io.StringIO()

        try:
            result = func(*args, **kwargs)
            output = captured_output.getvalue()
            return result, output
        finally:
            sys.stdout = old_stdout
            # Print captured output
            print(output, end='')

    return wrapper

def main():
    # Capture output
    old_stdout = sys.stdout
    sys.stdout = captured_output = io.StringIO()

    try:
        print("="*80)
        print("MODULE PERFORMANCE ANALYZER - First Principles Optimization")
        print("="*80)

        # Load training data
        df = load_training_data()
        if df is None:
            return

        # Analyze false positives
        fp_result = analyze_false_positives(df)
        problem_modules = fp_result[0] if fp_result else []

        # Analyze false negatives
        fn_result = analyze_false_negatives(df)
        weak_modules = fn_result[0] if fn_result else []

        # Generate recommendations
        recommendations = generate_optimization_recommendations(problem_modules, weak_modules, df)

        # Export
        export_recommendations(recommendations, Path("optimization_recommendations.txt"))

        print(f"\n{'='*80}")
        print("ANALYSIS COMPLETE")
        print(f"{'='*80}")

        # Get captured output
        report_content = captured_output.getvalue()

    finally:
        sys.stdout = old_stdout
        # Print captured output to console
        print(report_content, end='')

    # Generate timestamp
    timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Generate Word and PDF reports
    print(f"\n{'='*80}")
    print("GENERATING REPORTS")
    print(f"{'='*80}")

    word_path = generate_word_report(report_content, timestamp_str)
    pdf_path = generate_pdf_report(report_content, timestamp_str)

    # Summary
    print(f"\n{'='*80}")
    print("REPORT GENERATION COMPLETE")
    print(f"{'='*80}")
    if word_path:
        print(f"Word Report: {word_path}")
    if pdf_path:
        print(f"PDF Report: {pdf_path}")
    print(f"\nReports saved to: {REPORT_DIR}")

if __name__ == "__main__":
    main()
